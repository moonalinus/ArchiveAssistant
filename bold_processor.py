from docx import Document

def process_document(input_file, output_file):
    doc = Document(input_file)
    tables = doc.tables

    # Пропускаем первые 3 страницы
    if len(doc.sections) > 0:
        section = doc.sections[0]
        section.start_type = 3

    # Обрабатываем таблицы, начиная с четвертой страницы
    for table in tables:
        # Пропускаем первую строку (заголовок)
        for i in range(1, len(table.rows)):
            cell = table.cell(i, 2)  # Третий столбец (индекс 2)
            paragraphs = cell.paragraphs
            if len(paragraphs) >= 2:
                run = paragraphs[1].runs[0]
                run.bold = True

    doc.save(output_file)

# Пример использования
input_file = 'input.docx'
output_file = 'output.docx'
process_document(input_file, output_file)